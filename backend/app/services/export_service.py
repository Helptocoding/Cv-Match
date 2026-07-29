from io import BytesIO
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Flowable
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from xhtml2pdf import pisa

from app.models.export_schema import AdaptedCV
from app.services.template_service import TemplateService
from app.utils import cv_renderer
from app.utils.fonts import DEFAULT_FONT, resolve_font


class ExportService:
    def __init__(self) -> None:
        self.template_service = TemplateService()

    def build_harvard_pdf(self, adapted_cv: AdaptedCV, font_key: str = DEFAULT_FONT) -> bytes:
        try:
            return cv_renderer.render_cv_to_pdf(adapted_cv, font_key=font_key)
        except RuntimeError:
            pass

        html = self.template_service.render_harvard_html(adapted_cv)
        html_pdf = self._build_pdf_from_html(html)
        if html_pdf is not None:
            return html_pdf
        return self._build_reportlab_fallback(adapted_cv)

    def build_preview_html(self, adapted_cv: AdaptedCV, font_key: str = DEFAULT_FONT) -> str:
        return cv_renderer.render_cv_to_html(adapted_cv, font_key=font_key)

    def build_harvard_docx(self, adapted_cv: AdaptedCV) -> bytes:
        document = Document()
        source_cv = adapted_cv.source_cv

        experience_map = {}
        if source_cv:
            experience_map = {
                f"{exp.company.strip().lower()}||{exp.title.strip().lower()}": exp
                for exp in source_cv.experience
            }

        document.add_heading(source_cv.basics.full_name if source_cv else "Candidato", level=0)
        if source_cv and source_cv.basics.headline:
            document.add_paragraph(source_cv.basics.headline)

        if adapted_cv.adapted_summary:
            document.add_heading("Resumen Profesional", level=1)
            document.add_paragraph(adapted_cv.adapted_summary)

        document.add_heading("Experiencia", level=1)
        for experience in adapted_cv.adapted_experience:
            key = f"{experience.company.strip().lower()}||{experience.title.strip().lower()}"
            src = experience_map.get(key)
            dates = ""
            if src:
                if src.start_date and src.end_date:
                    dates = f"{src.start_date} – {src.end_date}"
                elif src.start_date:
                    dates = f"{src.start_date} – Presente"

            heading = " – ".join(part for part in [experience.title, experience.company] if part)
            paragraph = document.add_paragraph()
            paragraph.add_run(heading or "Experiencia").bold = True
            if dates:
                paragraph.add_run(f"  |  {dates}")
            if src and src.location:
                document.add_paragraph(src.location)
            for bullet in experience.rewritten_bullets:
                document.add_paragraph(bullet, style="List Bullet")

        if source_cv and source_cv.education:
            document.add_heading("Educación", level=1)
            for item in source_cv.education:
                paragraph = document.add_paragraph()
                paragraph.add_run(item.degree or "Título").bold = True
                if item.institution:
                    paragraph.add_run(f" – {item.institution}")
                dates = ""
                if item.start_date and item.end_date:
                    dates = f"{item.start_date} – {item.end_date}"
                elif item.end_date:
                    dates = item.end_date
                if dates:
                    paragraph.add_run(f"  |  {dates}")
                if item.field_of_study:
                    document.add_paragraph(item.field_of_study)

        if source_cv and source_cv.skills:
            document.add_heading("Habilidades", level=1)
            document.add_paragraph(" · ".join(skill.name for skill in source_cv.skills))

        if source_cv and source_cv.certifications:
            document.add_heading("Certificaciones", level=1)
            for cert in source_cv.certifications:
                paragraph = document.add_paragraph()
                paragraph.add_run(cert.name).bold = True
                if cert.issuer:
                    paragraph.add_run(f" – {cert.issuer}")
                if cert.date:
                    paragraph.add_run(f"  |  {cert.date}")

        if source_cv and source_cv.languages:
            document.add_heading("Idiomas", level=1)
            lang_parts = [
                f"{lang.name} ({lang.level})" if lang.level else lang.name
                for lang in source_cv.languages
            ]
            document.add_paragraph(" · ".join(lang_parts))

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def build_markdown(self, adapted_cv: AdaptedCV) -> str:
        source_cv = adapted_cv.source_cv
        lines: list[str] = []

        lines.append(f"# {source_cv.basics.full_name if source_cv else 'Candidato'}")
        if source_cv and source_cv.basics.headline:
            lines.append(source_cv.basics.headline)

        contact_parts = []
        if source_cv:
            if source_cv.basics.email:
                contact_parts.append(source_cv.basics.email)
            if source_cv.basics.phone:
                contact_parts.append(source_cv.basics.phone)
            if source_cv.basics.location:
                contact_parts.append(source_cv.basics.location)
        if contact_parts:
            lines.append(" · ".join(contact_parts))
        lines.append("")

        if adapted_cv.adapted_summary:
            lines.append("## Resumen Profesional")
            lines.append(adapted_cv.adapted_summary)
            lines.append("")

        lines.append("## Experiencia")
        experience_map = {}
        if source_cv:
            experience_map = {
                f"{exp.company.strip().lower()}||{exp.title.strip().lower()}": exp
                for exp in source_cv.experience
            }

        for experience in adapted_cv.adapted_experience:
            key = f"{experience.company.strip().lower()}||{experience.title.strip().lower()}"
            src = experience_map.get(key)
            dates = ""
            if src:
                if src.start_date and src.end_date:
                    dates = f" | {src.start_date} – {src.end_date}"
                elif src.start_date:
                    dates = f" | {src.start_date} – Presente"
            heading = " – ".join(part for part in [experience.title, experience.company] if part)
            lines.append(f"### {heading or 'Experiencia'}{dates}")
            if src and src.location:
                lines.append(f"*{src.location}*")
            for bullet in experience.rewritten_bullets:
                lines.append(f"- {bullet}")
            lines.append("")

        if source_cv and source_cv.education:
            lines.append("## Educación")
            for item in source_cv.education:
                dates = ""
                if item.start_date and item.end_date:
                    dates = f" | {item.start_date} – {item.end_date}"
                elif item.end_date:
                    dates = f" | {item.end_date}"
                heading = " – ".join(part for part in [item.degree, item.institution] if part)
                lines.append(f"### {heading or 'Educación'}{dates}")
                if item.field_of_study:
                    lines.append(f"*{item.field_of_study}*")
                lines.append("")

        if source_cv and source_cv.skills:
            lines.append("## Habilidades")
            lines.append(" · ".join(skill.name for skill in source_cv.skills))
            lines.append("")

        if source_cv and source_cv.certifications:
            lines.append("## Certificaciones")
            for cert in source_cv.certifications:
                entry = cert.name
                if cert.issuer:
                    entry += f" – {cert.issuer}"
                if cert.date:
                    entry += f" | {cert.date}"
                lines.append(f"- {entry}")
            lines.append("")

        if source_cv and source_cv.languages:
            lines.append("## Idiomas")
            lang_parts = [
                f"{lang.name} ({lang.level})" if lang.level else lang.name
                for lang in source_cv.languages
            ]
            lines.append(" · ".join(lang_parts))
            lines.append("")

        return "\n".join(lines).strip() + "\n"

    def _build_pdf_from_html(self, html: str) -> bytes | None:
        try:
            buffer = BytesIO()
            document: Any = pisa.CreatePDF(src=html, dest=buffer, encoding="utf-8")
            if getattr(document, "err", 1):
                return None
            return buffer.getvalue()
        except Exception:
            return None

    def _build_reportlab_fallback(self, adapted_cv: AdaptedCV) -> bytes:
        buffer = BytesIO()
        document = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=54, rightMargin=54, topMargin=48, bottomMargin=48)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("HarvardTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=16, leading=20)
        section_style = ParagraphStyle("HarvardSection", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=13, spaceBefore=10)
        body_style = ParagraphStyle("HarvardBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=13)
        sub_style = ParagraphStyle("HarvardSub", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=9, leading=12, textColor=(0.4, 0.4, 0.4))

        story: list[Flowable] = []
        candidate_name = adapted_cv.source_cv.basics.full_name if adapted_cv.source_cv else "Candidato"
        story.append(Paragraph(candidate_name, title_style))

        source_cv = adapted_cv.source_cv
        if source_cv and source_cv.basics.headline:
            story.append(Paragraph(source_cv.basics.headline, sub_style))
        story.append(Spacer(1, 10))

        experience_map = {}
        if source_cv:
            experience_map = {
                f"{exp.company.strip().lower()}||{exp.title.strip().lower()}": exp
                for exp in source_cv.experience
            }

        if adapted_cv.adapted_summary:
            story.append(Paragraph("Resumen Profesional", section_style))
            story.append(Paragraph(adapted_cv.adapted_summary, body_style))

        story.append(Paragraph("Experiencia", section_style))
        for experience in adapted_cv.adapted_experience:
            key = f"{experience.company.strip().lower()}||{experience.title.strip().lower()}"
            src = experience_map.get(key)
            dates = ""
            if src:
                if src.start_date and src.end_date:
                    dates = f"  |  {src.start_date} – {src.end_date}"
                elif src.start_date:
                    dates = f"  |  {src.start_date} – Presente"
            heading = " – ".join(part for part in [experience.title, experience.company] if part)
            story.append(Paragraph(f"{heading or 'Experiencia'}{dates}", body_style))
            for bullet in experience.rewritten_bullets:
                story.append(Paragraph(f"• {bullet}", body_style))
            story.append(Spacer(1, 5))

        if source_cv and source_cv.education:
            story.append(Paragraph("Educación", section_style))
            for item in source_cv.education:
                dates = ""
                if item.start_date and item.end_date:
                    dates = f"  |  {item.start_date} – {item.end_date}"
                elif item.end_date:
                    dates = f"  |  {item.end_date}"
                heading = " – ".join(part for part in [item.degree, item.institution] if part)
                story.append(Paragraph(f"{heading or 'Educación'}{dates}", body_style))
                if item.field_of_study:
                    story.append(Paragraph(item.field_of_study, sub_style))
                story.append(Spacer(1, 4))

        if source_cv and source_cv.skills:
            story.append(Paragraph("Habilidades", section_style))
            story.append(Paragraph(" · ".join(skill.name for skill in source_cv.skills), body_style))

        document.build(list(story))
        return buffer.getvalue()
