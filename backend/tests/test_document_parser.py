from io import BytesIO

from docx import Document

from app.services.document_parser import extract_text_from_docx


def test_extract_text_from_docx() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Senior Backend Engineer")
    document.add_paragraph("FastAPI and PostgreSQL")
    document.save(buffer)

    text = extract_text_from_docx(buffer.getvalue())

    assert "Senior Backend Engineer" in text
    assert "FastAPI" in text
